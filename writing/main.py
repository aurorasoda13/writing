import psycopg2
from flask import Flask, render_template, request, redirect, session, send_file, jsonify
from datetime import timedelta
from docx import Document
from docx.shared import RGBColor, Pt
from io import BytesIO

# Connessione al database (le tue credenziali originali)
connection = psycopg2.connect(
    user="postgres.rlrnsmfvujxqhqxmgegt",
    password='IdnYnvCWIMzfZ9BZ',
    host="aws-0-eu-north-1.pooler.supabase.com",
    port=6543,
    database="postgres"
)


app = Flask(__name__)
app.secret_key = 'una-chiave-super-segreta'
app.permanent_session_lifetime = timedelta(minutes=30)


@app.route('/')
def index():
    if session.get("utente"):
        return render_template('homea.html', nome=session.get("utente"))
    else:    
        return render_template('home.html')

@app.route('/accesso')
def accesso():
    return render_template('accedi.html')

@app.route('/accedi', methods=['POST'])
def accedi():
    nome = request.form['nome']
    psw = request.form['psw']

    cursor = connection.cursor()
    cursor.execute("SELECT nome FROM utente WHERE nome = %s AND psw = %s", (nome, psw,))
    result = cursor.fetchone()
    cursor.close()

    if result is None:
        return render_template("accedi.html", errore="Credenziali errate")
    else:
        session["utente"] = nome
        return render_template("homea.html", nome=nome)

@app.route('/registrazione')
def registrazione():
    return render_template("registrazione.html")

@app.route('/registrati', methods=['POST'])
def registrati():
    nome = request.form['nome']
    psw = request.form['psw']
    email = request.form['email']

    cursor = connection.cursor()
    cursor.execute("SELECT nome FROM utente WHERE nome = %s", (nome,))
    result = cursor.fetchone()
    cursor.close()

    if result is None:
        cursor = connection.cursor()
        cursor.execute("INSERT INTO utente (nome, psw, email) VALUES (%s, %s, %s)", (nome, psw, email))
        connection.commit()
        cursor.close()

        session["utente"] = nome
        session.permanent = False

        return render_template("homea.html", nome=nome, errore="Registrazione effettuata con successo")
    else:
        return render_template("registrazione.html", errore="Nome utente già esistente")

@app.route('/logout')
def logout():
    session.clear()
    return redirect('/')

@app.route('/nuovolibro')
def nuovolibro():
    if not session.get("utente"):
        return render_template('accedi.html')
    else:
        cursor = connection.cursor()
        cursor.execute("SELECT nome FROM genere")
        result = cursor.fetchall()
        cursor.execute("SELECT nome FROM genere")
        tags = cursor.fetchall()
        cursor.close()
        return render_template('nuovolibro.html', result=result, tags=tags)

@app.route('/executelibro', methods=['POST'])
def executelibro():
    titolo = request.form.get('titolo')
    trama = request.form.get('trama')
    sinossi = request.form.get('sinossi')
    nome_genere = request.form.get('nome')
    tags = request.form.getlist('tags[]')
    tagn = request.form.get('tagn')

    if not titolo or not trama or not sinossi or not nome_genere:
        return "Tutti i campi sono obbligatori", 400

    nome_utente = session.get("utente")

    cursor=connection.cursor()
    sql="SELECT idgenere FROM genere WHERE nome = %s"
    cursor.execute(sql, (nome_genere,))
    row = cursor.fetchone()
    print(row)
    if not row:
        return "Genere non valido", 400
    idgenere = row[0]

    cursor.execute("""
        INSERT INTO libro (titolo, trama, sinossi, idgenere, idautore)
        VALUES (%s, %s, %s, %s, %s)
        RETURNING id
    """, (titolo, trama, sinossi, idgenere, nome_utente))
    idlibro = cursor.fetchone()[0]

    if tagn:
        cursor.execute("INSERT INTO genere(nome) VALUES(%s)", (tagn,))
        tags.append(tagn)

    for tag in tags:
        cursor.execute("SELECT idgenere FROM genere WHERE nome = %s", (tag,))
        row = cursor.fetchone()
        if row:
            idgenere = row[0]
            cursor.execute("INSERT INTO genere_libro (idgenere, idlibro) VALUES (%s, %s)", (idgenere, idlibro))

    connection.commit()

    return render_template("homea.html", nome=nome_utente, errore="Libro inserito correttamente!")

@app.route('/ituoilibri')
def ituoilibri():
    if not session.get("utente"):
        return render_template('accedi.html')

    sql = """SELECT l.id, l.titolo, l.trama, g.nome 
             FROM libro l 
             INNER JOIN genere g ON l.idgenere = g.idgenere 
             INNER JOIN utente a ON l.idautore = a.nome 
             WHERE a.nome = %s"""
    cursor = connection.cursor()
    cursor.execute(sql, (session.get("utente"),))
    result = cursor.fetchall()
    return render_template("ituoilibri.html", result=result)

@app.route('/libro/<int:book_id>')
def libro(book_id):
    if not session.get("utente"):
        return render_template('accedi.html')
    
    cursor = connection.cursor()
    sql = """SELECT l.id, l.titolo, l.trama, l.sinossi, g.nome 
             FROM libro l 
             INNER JOIN genere g ON l.idgenere = g.idgenere 
             INNER JOIN utente a ON l.idautore = a.nome 
             WHERE a.nome = %s AND l.id = %s"""
    cursor.execute(sql, (session.get("utente"), book_id,))
    result = cursor.fetchone()

    sql = """SELECT genere.nome 
             FROM genere 
             INNER JOIN genere_libro ON genere.idgenere = genere_libro.idgenere 
             INNER JOIN libro ON libro.id = genere_libro.idlibro 
             WHERE libro.id = %s"""
    cursor.execute(sql, (book_id,))
    tags = cursor.fetchall()

    sql = """SELECT capitolo.titolo, capitolo.id
             FROM capitolo 
             INNER JOIN libro ON capitolo.idlibro = libro.id 
             INNER JOIN utente ON libro.idautore = utente.nome 
             WHERE idlibro = %s AND idautore = %s
             ORDER BY numero_cap"""
    cursor.execute(sql, (book_id, session.get("utente"),))
    capitoli = cursor.fetchall()
    cursor.close()

    return render_template("libro.html", result=result, tags=tags, capitoli=capitoli)

@app.route('/personaggio/<int:book_id>')
def personaggio(book_id):
    cursor = connection.cursor()
    cursor.execute("SELECT nome FROM ruolo")
    result = cursor.fetchall()
    cursor.close()
    return render_template('personaggio.html', result=result, idlibro=book_id)

@app.route('/executepersonaggio', methods=['POST'])
def executepersonaggio():
    nome = request.form['nome']
    alias = request.form['alias']
    descrizione_fisica = request.form['descrizione-fisica']
    ruolo = request.form['ruolo']
    psicologia = request.form['psicologia']
    obiettivi = request.form['obiettivi']
    background = request.form['background']
    note = request.form['note']
    book_id = request.form['idlibro']

    cursor = connection.cursor()
    cursor.execute("SELECT id FROM ruolo WHERE nome = %s", (ruolo,))
    idruolo = cursor.fetchone()[0]

    cursor.execute("""INSERT INTO personaggio 
                      (nome, alias, descrizione_fisica, idruolo, psicologia, obiettivi, background, note, idlibro) 
                      VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                   (nome, alias, descrizione_fisica, idruolo, psicologia, obiettivi, background, note, book_id,))
    
    connection.commit()
    cursor.close()
    return redirect("/ituoilibri")

@app.route('/capitolo/<int:book_id>')
def capitolo(book_id):
    cursor = connection.cursor()
    cursor.execute("""SELECT personaggio.nome 
                      FROM personaggio 
                      INNER JOIN libro ON personaggio.idlibro = libro.id 
                      INNER JOIN utente ON libro.idautore = utente.nome 
                      WHERE libro.id = %s AND libro.idautore = %s""", (book_id, session.get("utente"),))
    personaggi = cursor.fetchall()

    cursor.execute("""SELECT oggetto.nome 
                      FROM oggetto 
                      INNER JOIN libro ON oggetto.idlibro = libro.id  
                      INNER JOIN utente ON libro.idautore = utente.nome 
                      WHERE libro.id = %s AND libro.idautore = %s""", (book_id, session.get("utente"),))
    oggetti = cursor.fetchall()

    cursor.execute("""SELECT luogo.nome 
                      FROM luogo 
                      INNER JOIN libro ON luogo.idlibro = libro.id  
                      INNER JOIN utente ON libro.idautore = utente.nome 
                      WHERE libro.id = %s AND libro.idautore = %s""", (book_id, session.get("utente"),))
    luoghi = cursor.fetchall()
    cursor.close()

    return render_template("nuovocap.html", personaggi=personaggi, oggetti=oggetti, luoghi=luoghi, book_id=book_id)



@app.route('/executecap', methods=['POST'])
def executecap():
    titolo = request.form['titolo']
    numero_cap = request.form['capitolo']
    personaggi = request.form.getlist('personaggi[]')
    luoghi = request.form.getlist('luoghi[]')
    oggetti = request.form.getlist('oggetti[]')
    riassunto = request.form['riassunto']
    testo = request.form['testo']
    idlibro = request.form['book_id']

    try:
        with connection.cursor() as cursor:
            # Insert the chapter and get its ID
            sql = "INSERT INTO capitolo (titolo, numero_cap, riassunto, testo, idlibro) VALUES (%s, %s, %s, %s, %s) RETURNING id"
            cursor.execute(sql, (titolo, numero_cap, riassunto, testo, idlibro,))
            idcapitolo_result = cursor.fetchone()
            

            if idcapitolo_result is None:
                # Handle case where chapter insertion failed
                connection.rollback() # Rollback the transaction
                return "Errore nell'inserimento del capitolo.", 500

            idcapitolo = idcapitolo_result[0]

            # Insert associated characters
            for personaggio in personaggi:
                cursor.execute("SELECT id FROM personaggio WHERE nome = %s", (personaggio,))
                row = cursor.fetchone()
                
                if row:
                    idpersonaggio = row[0]
                    cursor.execute("INSERT INTO personaggio_capitolo (idpersonaggio, idcapitolo) VALUES (%s, %s)", (idpersonaggio, idcapitolo))

            # Insert associated locations
            for luogo in luoghi:
                cursor.execute("SELECT id FROM luogo WHERE nome = %s", (luogo,))
                row = cursor.fetchone()
                
                if row:
                    idluogo = row[0]
                    cursor.execute("INSERT INTO luogo_capitolo (idluogo, idcapitolo) VALUES (%s, %s)", (idluogo, idcapitolo))

            # Insert associated objects
            for oggetto in oggetti:
                cursor.execute("SELECT id FROM oggetto WHERE nome = %s", (oggetto,))
                row = cursor.fetchone()
                
                if row:
                    idoggetto = row[0]
                    cursor.execute("INSERT INTO oggetto_capitolo (idoggetto, idcapitolo) VALUES (%s, %s)", (idoggetto, idcapitolo))

        connection.commit() # Commit all changes at once
        return redirect(f'/libro/{idlibro}') # Redirect back to the book details page
    except Exception as e:
        connection.rollback() # Rollback on any error
        # Log the error for debugging purposes (you can use a logging library)
        print(f"Error inserting chapter: {e}")
        return "Si è verificato un errore durante l'inserimento del capitolo.", 500


@app.route('/luogo/<int:book_id>')
def luogo(book_id):
    sql = "SELECT nome FROM tipo_luogo"
    cursor = connection.cursor()
    cursor.execute(sql)
    tipo_luogo = cursor.fetchall()
    cursor.close()
    return render_template("luogo.html", tipo_luogo=tipo_luogo, book_id=book_id)

@app.route('/executeluogo', methods=['POST'])
def executeluogo():
    nome = request.form['nome']
    tipo_luogo = request.form['tipo']
    nuovo_tipo = request.form['luogon'].strip()
    descrizione = request.form['descrizione']
    eventiaccaduti = request.form['eventi']
    note = request.form['note']
    book_id = request.form['book_id']

    cursor = connection.cursor()

    try:
        # Se è stato inserito un nuovo tipo di luogo
        if nuovo_tipo:
            # Verifica se già esiste per evitare violazioni della chiave univoca
            cursor.execute("SELECT id FROM tipo_luogo WHERE nome = %s", (nuovo_tipo,))
            row = cursor.fetchone()
            if row:
                idtipo = row[0]
            else:
                cursor.execute("INSERT INTO tipo_luogo (nome) VALUES (%s) RETURNING id", (nuovo_tipo,))
                idtipo = cursor.fetchone()[0]
        else:
            cursor.execute("SELECT id FROM tipo_luogo WHERE nome = %s", (tipo_luogo,))
            row = cursor.fetchone()
            if row is None:
                cursor.close()
                return "Tipo luogo non valido", 400
            idtipo = row[0]

        # Inserimento luogo
        cursor.execute("""INSERT INTO luogo (nome, descrizione, eventiaccaduti, note, idlibro, idtipo)
                         VALUES (%s, %s, %s, %s, %s, %s) RETURNING id""",
                       (nome, descrizione, eventiaccaduti, note, book_id, idtipo,))
        cursor.fetchone()

        connection.commit()
        return redirect("/ituoilibri")

    except Exception as e:
        connection.rollback()
        print("Errore:", e)
        return "Errore durante l'inserimento del luogo.", 500
    finally:
        cursor.close()

@app.route('/oggetto/<int:book_id>')
def oggetto(book_id):
    sql="select nome from personaggio where idlibro=%s"
    cursor=connection.cursor()
    cursor.execute(sql, (book_id,))
    personaggi=cursor.fetchall()
    return render_template("oggetto.html",personaggi=personaggi, book_id=book_id)


@app.route('/executeoggetto', methods=['POST'])
def executeoggetto():
    nome = request.form['nome']
    proprietario = request.form['proprietario']
 
    descrizione = request.form['descrizione']
    funzione = request.form['funzione']
    potere = request.form['potere']
    storia = request.form['storia']
    provenienza = request.form['provenienza']
    book_id = request.form['book_id']

    cursor = connection.cursor()

    cursor.execute("SELECT id FROM personaggio WHERE nome = %s", (proprietario,))
    row = cursor.fetchone()
    if row is None:
        cursor.close()
        return "Personaggio non valido", 400
    idproprietario = row[0]

    cursor.execute("""INSERT INTO oggetto (nome, propretario ,descrizione, funzione, potere, storia, provenienza, idlibro)
                     VALUES (%s, %s, %s, %s, %s, %s, %s, %s)""",
                   (nome, idproprietario, descrizione, funzione, potere, storia, provenienza, book_id,))

    connection.commit()
    cursor.close()

    return redirect("/ituoilibri")

@app.route('/modificacap/<int:book_id>/<int:cap_id>')
def modificacap(book_id, cap_id):
    cursor = connection.cursor()

    # Capitolo details
    sql = "SELECT titolo, numero_cap, riassunto, testo FROM capitolo WHERE id = %s"
    cursor.execute(sql, (cap_id,))
    capitolo = cursor.fetchone()

    # Personaggi
    # Selected characters for the current chapter
    sql = """SELECT DISTINCT P.nome  
             FROM personaggio P 
             INNER JOIN personaggio_capitolo PC ON P.id = PC.idpersonaggio 
             WHERE PC.idcapitolo = %s"""
    cursor.execute(sql, (cap_id,))
    pers_sel = cursor.fetchall()

    # Other characters for the same book (not selected for this chapter)
    sql = """SELECT DISTINCT P.nome  
             FROM personaggio P 
             WHERE P.idlibro = %s 
             AND NOT EXISTS (
                 SELECT 1 
                 FROM personaggio_capitolo PC 
                 WHERE PC.idpersonaggio = P.id AND PC.idcapitolo = %s
             )"""
    cursor.execute(sql, (book_id, cap_id,))
    altri_pers = cursor.fetchall()


    # Luoghi
    # Selected locations for the current chapter
    sql = """SELECT DISTINCT L.nome  
             FROM luogo L 
             INNER JOIN luogo_capitolo LC ON L.id = LC.idluogo 
             WHERE LC.idcapitolo = %s"""
    cursor.execute(sql, (cap_id,))
    luoghi_sel = cursor.fetchall()

    # Other locations for the same book (not selected for this chapter)
    sql = """SELECT DISTINCT L.nome  
             FROM luogo L 
             WHERE L.idlibro = %s 
             AND NOT EXISTS (
                 SELECT 1 
                 FROM luogo_capitolo LC 
                 WHERE LC.idluogo = L.id AND LC.idcapitolo = %s
             )"""
    cursor.execute(sql, (book_id, cap_id,))
    altri_luoghi = cursor.fetchall()


    # Oggetti
    # Selected objects for the current chapter
    sql = """SELECT DISTINCT O.nome  
             FROM oggetto O 
             INNER JOIN oggetto_capitolo OC ON O.id = OC.idoggetto 
             WHERE OC.idcapitolo = %s"""
    cursor.execute(sql, (cap_id,))
    oggetti_sel = cursor.fetchall()

    # Other objects for the same book (not selected for this chapter)
    sql = """SELECT DISTINCT O.nome  
             FROM oggetto O 
             WHERE O.idlibro = %s 
             AND NOT EXISTS (
                 SELECT 1 
                 FROM oggetto_capitolo OC 
                 WHERE OC.idoggetto = O.id AND OC.idcapitolo = %s
             )"""
    cursor.execute(sql, (book_id, cap_id,))
    altri_oggetti = cursor.fetchall()

    cursor.close()

    return render_template("modifica_cap.html", 
                           capitolo=capitolo, 
                           luoghi_sel=luoghi_sel, 
                           altri_luoghi=altri_luoghi, 
                           pers_sel=pers_sel, 
                           altri_pers=altri_pers, 
                           oggetti_sel=oggetti_sel, 
                           altri_oggetti=altri_oggetti, 
                           book_id=book_id, 
                           cap_id=cap_id)


@app.route("/executecapm", methods=['POST'])
def executecapm():
    titolo = request.form['titolo']
    numero_cap = request.form['numero_cap'] 
    personaggi = request.form.getlist('personaggi[]')
    luoghi = request.form.getlist('luoghi[]')
    oggetti = request.form.getlist('oggetti[]')
    riassunto = request.form['riassunto']
    testo = request.form['testo']
    book_id = request.form['book_id']
    cap_id = request.form['cap_id'] 

    try:
        with connection.cursor() as cursor:
            # Update the chapter's main details
            sql = """UPDATE capitolo
                     SET titolo = %s, numero_cap = %s, riassunto = %s, testo = %s
                     WHERE id = %s;"""
            cursor.execute(sql, (titolo, numero_cap, riassunto, testo, cap_id,))

            # Delete existing associations for this chapter *once* before inserting new ones
            cursor.execute("DELETE FROM personaggio_capitolo WHERE idcapitolo = %s;", (cap_id,))
            cursor.execute("DELETE FROM luogo_capitolo WHERE idcapitolo = %s;", (cap_id,))
            cursor.execute("DELETE FROM oggetto_capitolo WHERE idcapitolo = %s;", (cap_id,))

            # Insert associated characters
            for personaggio in personaggi:
                cursor.execute("SELECT id FROM personaggio WHERE nome = %s", (personaggio,))
                row = cursor.fetchone()
                if row:
                    idpersonaggio = row[0]
                    cursor.execute("INSERT INTO personaggio_capitolo (idpersonaggio, idcapitolo) VALUES (%s, %s)", (idpersonaggio, cap_id))

            # Insert associated locations
            for luogo in luoghi:
                cursor.execute("SELECT id FROM luogo WHERE nome = %s", (luogo,))
                row = cursor.fetchone()
                if row:
                    idluogo = row[0]
                    cursor.execute("INSERT INTO luogo_capitolo (idluogo, idcapitolo) VALUES (%s, %s)", (idluogo, cap_id))

            # Insert associated objects
            for oggetto in oggetti:
                cursor.execute("SELECT id FROM oggetto WHERE nome = %s", (oggetto,))
                row = cursor.fetchone()
                if row:
                    idoggetto = row[0]
                    cursor.execute("INSERT INTO oggetto_capitolo (idoggetto, idcapitolo) VALUES (%s, %s)", (idoggetto, cap_id))

        connection.commit()  
        return redirect(f'/libro/{book_id}')
    except Exception as e:
        connection.rollback()  
        print(f"Error updating chapter: {e}")
        return "Si è verificato un errore durante l'aggiornamento del capitolo.", 500

@app.route('/download_book_docx/<int:book_id>')
def download_book_docx(book_id):
    if not session.get("utente"):
        return redirect('/accesso') # Assicurati che l'utente sia loggato

    try:
        doc = Document()

        # Imposta il colore per i titoli di livello 1 e 2 su nero
        # Questo assicura che i titoli del libro e dei capitoli siano neri per default.
        style_heading_1 = doc.styles['Heading 1']
        style_heading_1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        # style_heading_1.font.size = Pt(24) # Esempio: imposta la dimensione del font per Heading 1

        style_heading_2 = doc.styles['Heading 2']
        style_heading_2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        # style_heading_2.font.size = Pt(18) # Esempio: imposta la dimensione del font per Heading 2

        with connection.cursor() as cursor:
            # Recupera il titolo del libro
            sql_book = """
                SELECT titolo
                FROM libro
                WHERE id = %s AND idautore = %s
            """
            cursor.execute(sql_book, (book_id, session.get("utente"),))
            book_title_result = cursor.fetchone()

            if not book_title_result:
                return "Libro non trovato o non autorizzato.", 404

            book_title = book_title_result[0]

            # Aggiungi il titolo principale del documento (titolo del libro)
            doc.add_heading(book_title, level=1) 

            # Recupera i capitoli (titolo, numero, testo) ordinati per numero di capitolo
            sql_chapters = """
                SELECT titolo, numero_cap, testo
                FROM capitolo
                WHERE idlibro = %s
                ORDER BY numero_cap
            """
            cursor.execute(sql_chapters, (book_id,))
            chapters = cursor.fetchall()

            if not chapters:
                doc.add_paragraph("Nessun capitolo trovato per questo libro.")

            # Itera sui capitoli e aggiungili al documento
            for i, (chapter_title, chapter_num, chapter_text) in enumerate(chapters):
                # Aggiungi un'interruzione di pagina prima di ogni capitolo, tranne il primo
                if i > 0:
                    doc.add_page_break()
                
                doc.add_heading(f'Capitolo {chapter_num}: {chapter_title}', level=2)
                
                if chapter_text:
                    # Sostituisci i doppi a capo (o Windows style CRLF) con un a capo singolo
                    # per una formattazione migliore in Word, garantendo che i paragrafi
                    # siano gestiti correttamente.
                    formatted_text = chapter_text.replace('\r\n\r\n', '\r\n').replace('\n\n', '\n')
                    
                    # Split the text by single newlines to create separate paragraphs
                    # This helps in maintaining proper paragraph breaks in Word
                    paragraphs = formatted_text.split('\n')
                    for p_text in paragraphs:
                        doc.add_paragraph(p_text.strip()) # Aggiungi ogni riga come un nuovo paragrafo
                
        # Salva il documento in un buffer di byte in memoria
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0) # Riporta il "puntatore" del buffer all'inizio

        # Invia il file per il download
        filename = f"{book_title.replace(' ', '_').lower()}.docx"
        return send_file(bio, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    except Exception as e:
        # Rollback in caso di errore nel database
        connection.rollback() 
        print(f"Errore durante la generazione del documento Word: {e}")
        return "Si è verificato un errore durante la generazione del documento Word.", 500


@app.route('/eliminacapitolo', methods=['POST'])
def eliminacapitolo():
    if not session.get("utente"):
        return render_template('accedi.html')
    
    try:
        capitolo_id = request.form['capitolo_id']
        book_id = request.form['book_id']

        cursor = connection.cursor()
        cursor.execute("DELETE FROM capitolo WHERE id = %s AND idlibro = %s", (capitolo_id, book_id))
        connection.commit()
        cursor.close()

        return redirect(f"/libro/{book_id}")

    except Exception as e:
        connection.rollback()
        print("Errore durante l'eliminazione del capitolo:", e)
        return "Errore durante l'eliminazione del capitolo", 500

# Route API per recuperare personaggi di un libro (filtrati solo per idlibro)
@app.route('/api/personaggi/<int:book_id>')
def get_personaggi_by_book(book_id):
    if not session.get("utente"):
        return jsonify({"error": "Unauthorized"}), 401

    cursor = connection.cursor()
    # Query semplificata per filtrare solo per idlibro
    sql = """SELECT id, nome 
             FROM personaggio 
             WHERE idlibro = %s 
             ORDER BY nome"""
    cursor.execute(sql, (book_id,)) # Rimosso session.get("utente") dai parametri
    personaggi = cursor.fetchall()
    cursor.close()
    
    personaggi_list = [{"id": p[0], "nome": p[1]} for p in personaggi]
    return jsonify(personaggi_list)

# Route API per recuperare luoghi di un libro (filtrati solo per idlibro)
@app.route('/api/luoghi/<int:book_id>')
def get_luoghi_by_book(book_id):
    if not session.get("utente"):
        return jsonify({"error": "Unauthorized"}), 401

    cursor = connection.cursor()
    # Query semplificata per filtrare solo per idlibro
    sql = """SELECT id, nome 
             FROM luogo 
             WHERE idlibro = %s 
             ORDER BY nome"""
    cursor.execute(sql, (book_id,)) # Rimosso session.get("utente") dai parametri
    luoghi = cursor.fetchall()
    cursor.close()
    
    luoghi_list = [{"id": l[0], "nome": l[1]} for l in luoghi]
    return jsonify(luoghi_list)

# Route API per recuperare oggetti di un libro (filtrati solo per idlibro)
@app.route('/api/oggetti/<int:book_id>')
def get_oggetti_by_book(book_id):
    if not session.get("utente"):
        return jsonify({"error": "Unauthorized"}), 401

    cursor = connection.cursor()
    # Query semplificata per filtrare solo per idlibro
    sql = """SELECT id, nome 
             FROM oggetto 
             WHERE idlibro = %s 
             ORDER BY nome"""
    cursor.execute(sql, (book_id,)) # Rimosso session.get("utente") dai parametri
    oggetti = cursor.fetchall()
    cursor.close()
    
    oggetti_list = [{"id": o[0], "nome": o[1]} for o in oggetti]
    return jsonify(oggetti_list)

# Per visualizzare i dettagli dell'elemento selezionato (filtrati solo per ID elemento)
@app.route('/api/dettagli/<string:tipo>/<int:item_id>')
def get_item_details(tipo, item_id):
    if not session.get("utente"):
        return jsonify({"error": "Unauthorized"}), 401
    
    cursor = connection.cursor()
    print(item_id)
    details = {}
    print(tipo)
    try:
        if tipo == 'personagg':
            # Query semplificata per i dettagli del personaggio
            sql = """SELECT p.nome, p.alias, p.descrizione_fisica, p.psicologia, p.obiettivi, p.background, p.note, r.nome as ruolo_nome 
                     FROM personaggio p 
                     INNER JOIN ruolo r ON p.idruolo = r.id
                     WHERE p.id = %s""" # Rimosso filtro per idautore
            cursor.execute(sql, (item_id,)) # Rimosso session.get("utente") dai parametri
            result = cursor.fetchone()
            if result:
                details = {
                    "Nome": result[0],
                    "Alias": result[1],
                    "Descrizione Fisica": result[2],
                    "Psicologia": result[3],
                    "Obiettivi": result[4],
                    "Background": result[5],
                    "Note": result[6],
                    "Ruolo": result[7]
                }
        elif tipo == 'luogh':
            # Query semplificata per i dettagli del luogo
            sql = """SELECT lu.nome, lu.descrizione, lu.eventiaccaduti, lu.note, tl.nome as tipo_nome 
                     FROM luogo lu 
                     INNER JOIN tipo_luogo tl ON lu.idtipo = tl.id
                     WHERE lu.id = %s""" # Rimosso filtro per idautore
            cursor.execute(sql, (item_id,)) # Rimosso session.get("utente") dai parametri
            result = cursor.fetchone()
            print(result)
            if result:
                details = {
                    "Nome": result[0],
                    "Descrizione": result[1],
                    "Eventi Accaduti": result[2],
                    "Note": result[3],
                    "Tipo di Luogo": result[4]
                }
        elif tipo == 'oggett':
            # Query semplificata per i dettagli dell'oggetto
            sql = """SELECT o.nome, o.descrizione, o.funzione, o.potere, o.storia, o.provenienza, pers.nome as proprietario_nome
                     FROM oggetto o 
                     INNER JOIN personaggio pers ON o.propretario = pers.id 
                     WHERE o.id = %s""" # Rimosso filtro per idautore
            cursor.execute(sql, (item_id,)) # Rimosso session.get("utente") dai parametri
            result = cursor.fetchone()
            if result:
                details = {
                    "Nome": result[0],
                    "Descrizione": result[1],
                    "Funzione": result[2],
                    "Potere": result[3],
                    "Storia": result[4],
                    "Provenienza": result[5],
                    "Proprietario": result[6]
                }
        else:
            return jsonify({"error": "Tipo non valido"}), 400

    except Exception as e:
        print(f"Error fetching item details: {e}")
        return jsonify({"error": "Errore nel recupero dei dettagli."}), 500
    finally:
        cursor.close()

    if not details:
        return jsonify({"error": "Dettagli non trovati."}), 404
        
    return jsonify(details)

# Aggiungi questa nuova route in main.py
@app.route('/modificapersonaggio/<int:book_id>/<int:personaggio_id>')
def modificapersonaggio(book_id, personaggio_id):
    if not session.get("utente"):
        return redirect('/accesso')
    
    cursor = connection.cursor()

    # Get book details
    sql = """SELECT l.id, l.titolo, l.trama, l.sinossi, g.nome 
             FROM libro l 
             INNER JOIN genere g ON l.idgenere = g.idgenere 
             INNER JOIN utente a ON l.idautore = a.nome 
             WHERE a.nome = %s AND l.id = %s"""
    cursor.execute(sql, (session.get("utente"), book_id,))
    result = cursor.fetchone()

    # Get book tags
    sql_tags = """SELECT genere.nome 
                  FROM genere 
                  INNER JOIN genere_libro ON genere.idgenere = genere_libro.idgenere 
                  INNER JOIN libro ON libro.id = genere_libro.idlibro 
                  WHERE libro.id = %s"""
    cursor.execute(sql_tags, (book_id,))
    tags = cursor.fetchall()

    # Get character details
    sql_personaggio = """
        SELECT p.nome, p.alias, p.descrizione_fisica, p.psicologia, p.obiettivi, p.background, p.note, r.nome as ruolo
        FROM personaggio p
        INNER JOIN ruolo r ON p.idruolo = r.id
        WHERE p.id = %s AND p.idlibro = %s
    """
    cursor.execute(sql_personaggio, (personaggio_id, book_id))
    personaggio = cursor.fetchone()

    # Get all available roles
    sql_ruoli = "SELECT nome FROM ruolo"
    cursor.execute(sql_ruoli)
    ruoli = cursor.fetchall()
    
    cursor.close()

    print(f"idlibro: {book_id}, idpersonaggio: {personaggio_id}")
    if not personaggio or not result:
        return "Personaggio o libro non trovato.", 404

    return render_template(
        'modifica_personaggio.html',
        personaggio=personaggio,
        ruoli=ruoli,
        idlibro=book_id,
        idpersonaggio=personaggio_id,
        result=result,
        tags=tags
    )

@app.route('/editpersonaggio', methods=['POST'])
def editpersonaggio():
    nome = request.form['nome']
    alias = request.form['alias']
    descrizione_fisica = request.form['descrizione-fisica']
    ruolo = request.form['ruolo']
    psicologia = request.form['psicologia']
    obiettivi = request.form['obiettivi']
    background = request.form['background']
    note = request.form['note']
    book_id = request.form['idlibro']
    personaggio_id = request.form['idpersonaggio']

    cursor = connection.cursor()

    # Get role ID
    cursor.execute("SELECT id FROM ruolo WHERE nome = %s", (ruolo,))
    idruolo = cursor.fetchone()
    
    if not idruolo:
        cursor.close()
        return "Ruolo non valido", 400

    idruolo = idruolo[0]

    # Update character details
    cursor.execute("""
        UPDATE personaggio 
        SET nome = %s, alias = %s, descrizione_fisica = %s, idruolo = %s, psicologia = %s, obiettivi = %s, background = %s, note = %s 
        WHERE id = %s AND idlibro = %s
    """, (nome, alias, descrizione_fisica, idruolo, psicologia, obiettivi, background, note, personaggio_id, book_id))

    connection.commit()
    cursor.close()

    return redirect(f"/libro/{book_id}")

@app.route('/modificaluogo/<int:book_id>/<int:luogo_id>')
def modificaluogo(book_id, luogo_id):
    if not session.get("utente"):
        return redirect('/accesso')
    
    cursor = connection.cursor()

    # Get book details
    sql = """SELECT l.id, l.titolo, l.trama, l.sinossi, g.nome 
             FROM libro l 
             INNER JOIN genere g ON l.idgenere = g.idgenere 
             INNER JOIN utente a ON l.idautore = a.nome 
             WHERE a.nome = %s AND l.id = %s"""
    cursor.execute(sql, (session.get("utente"), book_id,))
    result = cursor.fetchone()

    # Get book tags
    sql_tags = """SELECT genere.nome 
                  FROM genere 
                  INNER JOIN genere_libro ON genere.idgenere = genere_libro.idgenere 
                  INNER JOIN libro ON libro.id = genere_libro.idlibro 
                  WHERE libro.id = %s"""
    cursor.execute(sql_tags, (book_id,))
    tags = cursor.fetchall()

    # Get location details
    sql_luogo = """
        SELECT lu.nome, lu.descrizione, lu.eventiaccaduti, lu.note, tl.nome as tipo_nome, lu.id
        FROM luogo lu
        INNER JOIN tipo_luogo tl ON lu.idtipo = tl.id
        WHERE lu.id = %s AND lu.idlibro = %s
    """
    cursor.execute(sql_luogo, (luogo_id, book_id))
    luogo = cursor.fetchone()

    # Get all available types of places
    sql_tipo_luogo = "SELECT nome FROM tipo_luogo"
    cursor.execute(sql_tipo_luogo)
    tipi_luoghi = cursor.fetchall()
    
    cursor.close()

    if not luogo or not result:
        return "Luogo o libro non trovato.", 404

    return render_template(
        'modifica_luogo.html',
        luogo=luogo,
        tipi_luoghi=tipi_luoghi,
        idlibro=book_id,
        idluogo=luogo_id,
        result=result,
        tags=tags
    )
@app.route('/editluogo', methods=['POST'])
def editluogo():
    
    nome = request.form['nome']
    tipo_luogo = request.form['tipo']
    nuovo_tipo = request.form['luogon'].strip()
    descrizione = request.form['descrizione']
    eventiaccaduti = request.form['eventi']
    note = request.form['note']
    book_id = request.form['book_id']
    luogo_id = request.form['luogo_id']


    print(f"nome: {nome}")
    print(f"luogo_id: {luogo_id}")
    print(f"book_id: {book_id}")

    cursor = connection.cursor()

    try:
        # Se è stato inserito un nuovo tipo di luogo
        if nuovo_tipo:
            # Verifica se già esiste per evitare violazioni della chiave univoca
            cursor.execute("SELECT id FROM tipo_luogo WHERE nome = %s", (nuovo_tipo,))
            row = cursor.fetchone()
            if row:
                idtipo = row[0]
            else:
                cursor.execute("INSERT INTO tipo_luogo (nome) VALUES (%s) RETURNING id", (nuovo_tipo,))
                idtipo = cursor.fetchone()[0]
        else:
            cursor.execute("SELECT id FROM tipo_luogo WHERE nome = %s", (tipo_luogo,))
            row = cursor.fetchone()
            if row is None:
                cursor.close()
                return "Tipo luogo non valido", 400
            idtipo = row[0]

        # Update location details
        cursor.execute("""UPDATE luogo 
                          SET nome = %s, descrizione = %s, eventiaccaduti = %s, note = %s, idtipo = %s 
                          WHERE id = %s AND idlibro = %s""",
                       (nome, descrizione, eventiaccaduti, note, idtipo, luogo_id, book_id))

        connection.commit()
        return redirect(f"/libro/{book_id}")

    except Exception as e:
        connection.rollback()
        print("Errore:", e)
        return "Errore durante l'aggiornamento del luogo.", 500
    finally:
        cursor.close()


@app.route('/modificaoggetto/<int:book_id>/<int:oggetto_id>')
def modificaoggetto(book_id, oggetto_id):
    if not session.get("utente"):
        return render_template('accedi.html')

    cursor = connection.cursor()

    # Recupera i dati dell'oggetto
    cursor.execute("""
        SELECT o.id, o.nome, p.nome, o.descrizione, o.funzione, o.potere, o.storia, o.provenienza
        FROM oggetto o
        LEFT JOIN personaggio p ON o.propretario = p.id
        WHERE o.idlibro = %s AND o.id = %s
    """, (book_id, oggetto_id,))
    oggetto = cursor.fetchone()

    # Recupera tutti i personaggi per la lista a discesa del proprietario
    cursor.execute("""
        SELECT nome FROM personaggio WHERE idlibro = %s
    """, (book_id,))
    personaggi = cursor.fetchall()

    # Recupera i dati del libro per la sidebar
    sql = """SELECT l.id, l.titolo, l.trama, l.sinossi, g.nome
             FROM libro l
             INNER JOIN genere g ON l.idgenere = g.idgenere
             WHERE l.id = %s"""
    cursor.execute(sql, (book_id,))
    result = cursor.fetchone()

    sql = """SELECT genere.nome 
             FROM genere 
             INNER JOIN genere_libro ON genere.idgenere = genere_libro.idgenere 
             INNER JOIN libro ON libro.id = genere_libro.idlibro 
             WHERE libro.id = %s"""
    cursor.execute(sql, (book_id,))
    tags = cursor.fetchall()

    cursor.close()

    if oggetto is None:
        return "Oggetto non trovato", 404

    return render_template('modifica_oggetto.html', 
                            oggetto=oggetto, 
                            personaggi=personaggi, 
                            book_id=book_id,
                            result=result,
                            tags=tags)




@app.route('/editoggetto', methods=['POST'])
def editoggetto():
    if not session.get("utente"):
        return render_template('accedi.html')

    try:
        nome = request.form['nome']
        proprietario_nome = request.form.get('proprietario')
        descrizione = request.form['descrizione']
        funzione = request.form['funzione']
        potere = request.form['potere']
        storia = request.form['storia']
        provenienza = request.form['provenienza']
        book_id = request.form['book_id']
        oggetto_id = request.form['oggetto_id']

        cursor = connection.cursor()

        idproprietario = None
        if proprietario_nome:
            cursor.execute("SELECT id FROM personaggio WHERE nome = %s AND idlibro = %s", (proprietario_nome, book_id,))
            row = cursor.fetchone()
            if row:
                idproprietario = row[0]

        cursor.execute("""
            UPDATE oggetto
            SET nome = %s, propretario = %s, descrizione = %s, funzione = %s, potere = %s, storia = %s, provenienza = %s
            WHERE id = %s AND idlibro = %s
        """, (nome, idproprietario, descrizione, funzione, potere, storia, provenienza, oggetto_id, book_id,))

        connection.commit()
        cursor.close()
        return redirect(f"/libro/{book_id}")

    except Exception as e:
        connection.rollback()
        print("Errore durante l'aggiornamento dell'oggetto:", e)
        return "Errore durante l'aggiornamento dell'oggetto", 500

@app.route('/elimina_libro/<int:book_id>')
def elimina_libro(book_id):
    if not session.get("utente"):
        return redirect('/accesso')

    cursor = connection.cursor()

    try:
        # Elimina prima i tag associati al libro
        cursor.execute("DELETE FROM genere_libro WHERE idlibro = %s", (book_id,))
        
        # Poi elimina tutti i personaggi, capitoli, luoghi e oggetti
        # associati al libro per rispettare i vincoli di integrità referenziale
        cursor.execute("DELETE FROM personaggio WHERE idlibro = %s", (book_id,))
        cursor.execute("DELETE FROM capitolo WHERE idlibro = %s", (book_id,))
        cursor.execute("DELETE FROM luogo WHERE idlibro = %s", (book_id,))
        cursor.execute("DELETE FROM oggetto WHERE idlibro = %s", (book_id,))

        # Infine, elimina il libro
        cursor.execute("DELETE FROM libro WHERE id = %s", (book_id,))

        connection.commit()
        cursor.close()

        # Reindirizza l'utente alla homepage dopo la cancellazione
        return redirect('/')

    except Exception as e:
        connection.rollback()
        print("Errore durante l'eliminazione del libro:", e)
        return "Errore durante l'eliminazione del libro", 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=81)